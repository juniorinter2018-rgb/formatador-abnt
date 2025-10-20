# motor.py (Versão com parser de Delta para suportar formatação inline)

import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# (função adicionar_paragrafo_pre_textual não muda)
def adicionar_paragrafo_pre_textual(document, text, space_before=0, font_size=12, is_bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, is_upper=True):
    if text:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        run_text = text.upper() if is_upper else text
        run = p.add_run(run_text)
        run.font.size = Pt(font_size)
        run.bold = is_bold
        p.alignment = alignment

def aplicar_formatacao_paragrafo(p, attrs):
    """Aplica formatação a nível de parágrafo (headers, blockquote, etc.)"""
    if not attrs:
        # Formatação de parágrafo padrão
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return

    if attrs.get('header'):
        nivel = attrs.get('header')
        # Tenta aplicar o estilo ao primeiro 'run' (pedaço de texto)
        if p.runs:
            run = p.runs[0] 
            if nivel == 1:
                run.text = run.text.upper()
                run.bold = True
            if nivel == 2:
                run.bold = True
        # (Nota: a numeração de seção 1.1, 1.2 etc. foi temporariamente removida
        # para implementar esta correção e pode ser adicionada depois)
    
    elif attrs.get('blockquote'):
        fmt = p.paragraph_format
        fmt.left_indent = Cm(4)
        fmt.line_spacing = 1.0
        for run in p.runs: # Aplica fonte menor a todos os runs do parágrafo
            run.font.size = Pt(10)
    
    else:
        # Parágrafo padrão se nenhum atributo de bloco for encontrado
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        
def gerar_documento(info_trabalho, dados_delta, dados_referencias):
    document = docx.Document()

    # --- CONFIGURAÇÃO INICIAL DO DOCUMENTO ---
    style = document.styles['Normal']
    font = style.font; font.name = 'Arial'; font.size = Pt(12)
    for section in document.sections:
        section.top_margin = Cm(3); section.bottom_margin = Cm(2)
        section.left_margin = Cm(3); section.right_margin = Cm(2)

    # --- ELEMENTOS PRÉ-TEXTUAIS (Capa e Folha de Rosto) ---
    # (Esta seção não muda, pois usa info_trabalho)
    if info_trabalho.get('autor') and info_trabalho.get('titulo'):
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('instituicao'))
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('curso'), 12)
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('autor'), 120)
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('titulo'), 120, is_bold=True)
        if info_trabalho.get('subtitulo'):
            adicionar_paragrafo_pre_textual(document, info_trabalho.get('subtitulo'), 12)
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('cidade'), 150)
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('ano'), 12)
        document.add_page_break()
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('autor'), 0)
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('titulo'), 120, is_bold=True)
        if info_trabalho.get('subtitulo'):
            adicionar_paragrafo_pre_textual(document, info_trabalho.get('subtitulo'), 12)
        nota_curso = info_trabalho.get('curso', '[Nome do Curso]')
        nota_instituicao = info_trabalho.get('instituicao', '[Nome da Instituição]')
        nota = (f"Trabalho de Conclusão de Curso apresentado ao curso de {nota_curso} da {nota_instituicao}, "
                f"como requisito parcial para a obtenção do título de Bacharel.")
        p_nota = document.add_paragraph(); p_nota.paragraph_format.space_before = Pt(100)
        p_nota.paragraph_format.left_indent = Cm(8); p_nota.add_run(nota)
        if info_trabalho.get('orientador'):
            adicionar_paragrafo_pre_textual(document, f"Orientador(a): {info_trabalho.get('orientador')}", 24, alignment=WD_ALIGN_PARAGRAPH.LEFT, is_upper=False)
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('cidade'), 120)
        adicionar_paragrafo_pre_textual(document, info_trabalho.get('ano'), 12)
    
    # --- NOVA LÓGICA DE PROCESSAMENTO (Lendo o Delta) ---
    
    # (A lógica de Sumário foi desativada pois depende da lógica antiga.
    # Terá de ser reimplementada futuramente)
    
    # Adiciona o Corpo do Texto
    document.add_page_break()
    
    p = document.add_paragraph() # Inicia o primeiro parágrafo
    paragrafo_attrs = {} # Atributos do parágrafo atual

    if dados_delta and dados_delta.get('ops'):
        for op in dados_delta.get('ops'):
            if 'insert' in op:
                texto = op['insert']
                attrs = op.get('attributes', {})
                
                if '\n' in texto:
                    # O op contém quebra de linha, o que define um parágrafo
                    partes = texto.split('\n')
                    
                    # 1. Adiciona a primeira parte ao parágrafo atual
                    if partes[0]:
                        run = p.add_run(partes[0])
                        run.bold = attrs.get('bold', False)
                        run.italic = attrs.get('italic', False)
                    
                    # 2. Aplica a formatação de bloco ao parágrafo que acabamos de fechar
                    paragrafo_attrs = attrs # Atributos de bloco estão no '\n'
                    aplicar_formatacao_paragrafo(p, paragrafo_attrs)
                    
                    # 3. Cria novos parágrafos para as partes intermediárias (se houver)
                    for i in range(1, len(partes) - 1):
                        p = document.add_paragraph()
                        if partes[i]:
                            run = p.add_run(partes[i])
                            run.bold = attrs.get('bold', False)
                            run.italic = attrs.get('italic', False)
                        aplicar_formatacao_paragrafo(p, paragrafo_attrs)
                    
                    # 4. Inicia o *próximo* parágrafo
                    p = document.add_paragraph()
                    if partes[-1]: # Se houver texto após a última quebra de linha
                        run = p.add_run(partes[-1])
                        run.bold = attrs.get('bold', False)
                        run.italic = attrs.get('italic', False)

                else:
                    # O op é apenas texto inline (bold, italic, etc.)
                    run = p.add_run(texto)
                    run.bold = attrs.get('bold', False)
                    run.italic = attrs.get('italic', False)

    # Aplica formatação ao último parágrafo
    aplicar_formatacao_paragrafo(p, paragrafo_attrs)


    # --- Adiciona as Referências ---
    # (Esta seção não muda, pois usa dados_referencias)
    if dados_referencias:
        document.add_page_break()
        p_ref = document.add_paragraph('REFERÊNCIAS')
        if p_ref.runs: p_ref.runs[0].bold = True
        p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        dados_referencias.sort(key=lambda x: x.get('autor', ''))
        for ref in dados_referencias:
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            tipo = ref.get('tipo')
            if tipo == 'livro':
                p.add_run(f"{ref.get('autor', '')}. ")
                run_titulo = p.add_run(ref.get('titulo', '')); run_titulo.bold = True
                p.add_run(f". {ref.get('cidade', '')}: {ref.get('editora', '')}, {ref.get('ano', '')}.")
            elif tipo == 'site':
                p.add_run(f"{ref.get('autor', '')}. {ref.get('titulo', '')}. {ref.get('nome_site', '')}, {ref.get('ano', '')}. Disponível em: <{ref.get('url', '')}>. Acesso em: {ref.get('data_acesso', '')}.")
            elif tipo == 'artigo':
                p.add_run(f"{ref.get('autor', '')}. {ref.get('titulo', '')}. ")
                run_revista = p.add_run(ref.get('nome_revista', '')); run_revista.bold = True
                p.add_run(f", v. {ref.get('volume', '')}, n. {ref.get('numero', '')}, {ref.get('paginas', '')}, {ref.get('ano', '')}.")

    # --- SALVA O DOCUMENTO EM MEMÓRIA ---
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream